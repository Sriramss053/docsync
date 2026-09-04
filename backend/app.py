import os
import io
from datetime import datetime

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_socketio import SocketIO, join_room, emit
from dotenv import load_dotenv

from models import db, User, Document, Version, Comment, Collaborator
from auth import (
    login_required, admin_required, generate_token,
    require_doc_role, get_doc_role, ROLE_RANK,
)
from crypto_utils import encrypt_text, decrypt_text
from storage import save_file

load_dotenv()

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "dev-secret")
app.config["JWT_SECRET_KEY"] = os.environ.get("JWT_SECRET_KEY", "dev-jwt-secret")
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("DATABASE_URL", "sqlite:///docsync.db")
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

FRONTEND_ORIGIN = os.environ.get("FRONTEND_ORIGIN", "http://localhost:5173")
CORS(app, resources={r"/api/*": {"origins": FRONTEND_ORIGIN}}, supports_credentials=True)

db.init_app(app)
socketio = SocketIO(app, cors_allowed_origins=FRONTEND_ORIGIN, async_mode="threading")

with app.app_context():
    db.create_all()


# ---------------------------------------------------------------------------
# Health check
# ---------------------------------------------------------------------------
@app.route("/api/health")
def health():
    return jsonify({"status": "ok", "time": datetime.utcnow().isoformat()})


# ---------------------------------------------------------------------------
# Auth routes
# ---------------------------------------------------------------------------
@app.route("/api/auth/signup", methods=["POST"])
def signup():
    data = request.get_json(force=True)
    name = (data.get("name") or "").strip()
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""

    if not name or not email or not password:
        return jsonify({"error": "name, email and password are required"}), 400
    if len(password) < 6:
        return jsonify({"error": "password must be at least 6 characters"}), 400
    if User.query.filter_by(email=email).first():
        return jsonify({"error": "An account with this email already exists"}), 409

    # First user in the system becomes admin automatically
    is_first_user = User.query.count() == 0
    user = User(name=name, email=email, role="admin" if is_first_user else "editor")
    user.set_password(password)
    db.session.add(user)
    db.session.commit()

    token = generate_token(user)
    return jsonify({"token": token, "user": user.to_dict()}), 201


@app.route("/api/auth/login", methods=["POST"])
def login():
    data = request.get_json(force=True)
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""

    user = User.query.filter_by(email=email).first()
    if not user or not user.check_password(password):
        return jsonify({"error": "Invalid email or password"}), 401

    token = generate_token(user)
    return jsonify({"token": token, "user": user.to_dict()})


@app.route("/api/auth/me", methods=["GET"])
@login_required
def me():
    return jsonify({"user": request.current_user.to_dict()})


# ---------------------------------------------------------------------------
# Document CRUD
# ---------------------------------------------------------------------------
@app.route("/api/documents", methods=["GET"])
@login_required
def list_documents():
    user = request.current_user
    owned = Document.query.filter_by(owner_id=user.id)
    shared_ids = [c.doc_id for c in Collaborator.query.filter_by(user_id=user.id).all()]
    shared = Document.query.filter(Document.id.in_(shared_ids)) if shared_ids else []

    docs = {d.id: d for d in owned}
    for d in shared:
        docs[d.id] = d

    result = []
    for d in docs.values():
        item = d.to_dict()
        item["role"] = get_doc_role(user, d)
        result.append(item)
    result.sort(key=lambda x: x["updated_at"], reverse=True)
    return jsonify({"documents": result})


@app.route("/api/documents", methods=["POST"])
@login_required
def create_document():
    data = request.get_json(force=True) or {}
    title = data.get("title", "Untitled Document")
    folder = data.get("folder", "root")
    is_sensitive = bool(data.get("is_sensitive", False))

    doc = Document(
        title=title,
        owner_id=request.current_user.id,
        folder=folder,
        content_encrypted=encrypt_text(""),
        is_sensitive=is_sensitive,
    )
    db.session.add(doc)
    db.session.commit()
    return jsonify({"document": doc.to_dict()}), 201


@app.route("/api/documents/<doc_id>", methods=["GET"])
@login_required
@require_doc_role("viewer")
def get_document(doc_id):
    doc = request.doc
    item = doc.to_dict()
    item["content"] = decrypt_text(doc.content_encrypted)
    item["role"] = request.doc_role
    return jsonify({"document": item})


@app.route("/api/documents/<doc_id>", methods=["PUT"])
@login_required
@require_doc_role("editor")
def update_document(doc_id):
    doc = request.doc
    data = request.get_json(force=True) or {}

    if "title" in data:
        doc.title = data["title"]
    if "folder" in data:
        doc.folder = data["folder"]
    if "content" in data:
        # snapshot previous version before overwriting
        version = Version(
            doc_id=doc.id,
            content_encrypted=doc.content_encrypted,
            edited_by=request.current_user.id,
        )
        db.session.add(version)
        doc.content_encrypted = encrypt_text(data["content"])

    doc.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify({"document": doc.to_dict()})


@app.route("/api/documents/<doc_id>", methods=["DELETE"])
@login_required
@require_doc_role("admin")
def delete_document(doc_id):
    db.session.delete(request.doc)
    db.session.commit()
    return jsonify({"success": True})


# ---------------------------------------------------------------------------
# Collaborators / roles
# ---------------------------------------------------------------------------
@app.route("/api/documents/<doc_id>/collaborators", methods=["GET"])
@login_required
@require_doc_role("viewer")
def list_collaborators(doc_id):
    collabs = Collaborator.query.filter_by(doc_id=doc_id).all()
    result = []
    for c in collabs:
        u = User.query.get(c.user_id)
        result.append({"user_id": c.user_id, "name": u.name if u else "?", "email": u.email if u else "?", "role": c.role})
    return jsonify({"collaborators": result})


@app.route("/api/documents/<doc_id>/collaborators", methods=["POST"])
@login_required
@require_doc_role("admin")
def add_collaborator(doc_id):
    data = request.get_json(force=True) or {}
    email = (data.get("email") or "").strip().lower()
    role = data.get("role", "viewer")
    if role not in ROLE_RANK:
        return jsonify({"error": "invalid role"}), 400

    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify({"error": "No user found with that email"}), 404

    existing = Collaborator.query.filter_by(doc_id=doc_id, user_id=user.id).first()
    if existing:
        existing.role = role
    else:
        db.session.add(Collaborator(doc_id=doc_id, user_id=user.id, role=role))
    db.session.commit()
    return jsonify({"success": True})


# ---------------------------------------------------------------------------
# Version history
# ---------------------------------------------------------------------------
@app.route("/api/documents/<doc_id>/versions", methods=["GET"])
@login_required
@require_doc_role("viewer")
def list_versions(doc_id):
    versions = Version.query.filter_by(doc_id=doc_id).order_by(Version.timestamp.desc()).all()
    return jsonify({"versions": [v.to_dict() for v in versions]})


@app.route("/api/documents/<doc_id>/versions/<version_id>", methods=["GET"])
@login_required
@require_doc_role("viewer")
def get_version(doc_id, version_id):
    version = Version.query.filter_by(id=version_id, doc_id=doc_id).first()
    if not version:
        return jsonify({"error": "Version not found"}), 404
    item = version.to_dict()
    item["content"] = decrypt_text(version.content_encrypted)
    return jsonify({"version": item})


@app.route("/api/documents/<doc_id>/versions/<version_id>/restore", methods=["POST"])
@login_required
@require_doc_role("editor")
def restore_version(doc_id, version_id):
    doc = request.doc
    version = Version.query.filter_by(id=version_id, doc_id=doc_id).first()
    if not version:
        return jsonify({"error": "Version not found"}), 404

    # snapshot current content first, then restore
    snapshot = Version(doc_id=doc.id, content_encrypted=doc.content_encrypted, edited_by=request.current_user.id)
    db.session.add(snapshot)
    doc.content_encrypted = version.content_encrypted
    doc.updated_at = datetime.utcnow()
    db.session.commit()

    socketio.emit("document_restored", {"doc_id": doc_id, "content": decrypt_text(doc.content_encrypted)}, room=doc_id)
    return jsonify({"document": doc.to_dict()})


# ---------------------------------------------------------------------------
# Comments
# ---------------------------------------------------------------------------
@app.route("/api/documents/<doc_id>/comments", methods=["GET"])
@login_required
@require_doc_role("viewer")
def list_comments(doc_id):
    comments = Comment.query.filter_by(doc_id=doc_id).order_by(Comment.timestamp.asc()).all()
    return jsonify({"comments": [c.to_dict() for c in comments]})


@app.route("/api/documents/<doc_id>/comments", methods=["POST"])
@login_required
@require_doc_role("viewer")
def add_comment(doc_id):
    data = request.get_json(force=True) or {}
    text = (data.get("text") or "").strip()
    anchor = data.get("anchor")
    if not text:
        return jsonify({"error": "Comment text required"}), 400

    comment = Comment(doc_id=doc_id, user_id=request.current_user.id, text=text, anchor=anchor)
    db.session.add(comment)
    db.session.commit()

    payload = comment.to_dict()
    payload["user_name"] = request.current_user.name
    socketio.emit("new_comment", payload, room=doc_id)
    return jsonify({"comment": payload}), 201


# ---------------------------------------------------------------------------
# Export
# ---------------------------------------------------------------------------
@app.route("/api/documents/<doc_id>/export/<fmt>", methods=["GET"])
@login_required
@require_doc_role("viewer")
def export_document(doc_id, fmt):
    doc = request.doc
    content = decrypt_text(doc.content_encrypted)

    if fmt == "pdf":
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 8, txt=content or " ")
        buf = io.BytesIO(pdf.output(dest="S"))
        buf.seek(0)
        return send_file(buf, as_attachment=True, download_name=f"{doc.title}.pdf", mimetype="application/pdf")

    if fmt == "docx":
        from docx import Document as DocxDocument
        docx_doc = DocxDocument()
        docx_doc.add_heading(doc.title, level=1)
        for line in (content or "").split("\n"):
            docx_doc.add_paragraph(line)
        buf = io.BytesIO()
        docx_doc.save(buf)
        buf.seek(0)
        return send_file(
            buf, as_attachment=True, download_name=f"{doc.title}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    return jsonify({"error": "Unsupported format. Use 'pdf' or 'docx'."}), 400


# ---------------------------------------------------------------------------
# WebSocket events - real-time collaborative editing + chat
# ---------------------------------------------------------------------------
@socketio.on("join_document")
def on_join(data):
    doc_id = data.get("doc_id")
    user_name = data.get("user_name", "Someone")
    join_room(doc_id)
    emit("user_joined", {"user_name": user_name}, room=doc_id, include_self=False)


@socketio.on("leave_document")
def on_leave(data):
    doc_id = data.get("doc_id")
    user_name = data.get("user_name", "Someone")
    emit("user_left", {"user_name": user_name}, room=doc_id, include_self=False)


@socketio.on("content_change")
def on_content_change(data):
    """Broadcast live edits to everyone else in the document room.
    NOTE: uses last-write-wins broadcast sync (simple approach). For true
    conflict-free concurrent editing, swap this for a CRDT/OT library."""
    doc_id = data.get("doc_id")
    content = data.get("content")
    user_name = data.get("user_name", "Someone")
    emit("content_update", {"content": content, "user_name": user_name}, room=doc_id, include_self=False)


@socketio.on("cursor_move")
def on_cursor_move(data):
    doc_id = data.get("doc_id")
    emit("cursor_update", data, room=doc_id, include_self=False)


@socketio.on("chat_message")
def on_chat_message(data):
    doc_id = data.get("doc_id")
    emit("chat_message", data, room=doc_id)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    debug_mode = os.environ.get("FLASK_DEBUG", "0") == "1"
    socketio.run(app, host="0.0.0.0", port=port, debug=debug_mode, allow_unsafe_werkzeug=True)
